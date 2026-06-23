# Genera una versione .docx EDITABILE dell'Appendice B (partita d'esempio) dal manuale.
# Uso: python3 docs/genera_docx_appendiceB.py
# Estrae la sola Appendice B da docs/frenemies_manuale.md e la converte mantenendo la struttura:
#   - titoli  -> heading di Word
#   - righe d'azione ({: .azione })      -> paragrafo (con il grassetto dell'autore)
#   - box di fiction (<div class=fiction>) -> paragrafo rientrato e ombreggiato, didascalia in grassetto corsivo
#   - note "Le carte in mano" (> ...)    -> stile citazione, leggermente ombreggiato
import re
from docx import Document
from docx.shared import Pt, Inches
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

SRC = "docs/frenemies_manuale.md"
OUT = "docs/Appendice_B_partita_esempio.docx"

raw = open(SRC, encoding="utf-8").read().split("\n")
start = next(i for i, l in enumerate(raw) if l.startswith("# Appendice B"))
end = next(i for i, l in enumerate(raw) if i > start and l.startswith("# Appendice C"))
block = raw[start:end]

doc = Document()
doc.styles["Normal"].font.name = "Georgia"
doc.styles["Normal"].font.size = Pt(11)

def shade(par, fill):
    pPr = par._p.get_or_add_pPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear"); shd.set(qn("w:color"), "auto"); shd.set(qn("w:fill"), fill)
    pPr.append(shd)

def add_inline(par, text, bold=False, italic=False):
    # gestisce **grassetto**, <em>corsivo</em>, <b>grassetto</b>
    for tok in re.split(r"(\*\*|\*|<em>|</em>|<b>|</b>)", text):
        if tok == "**": bold = not bold; continue
        if tok == "*": italic = not italic; continue
        if tok in ("<em>",): italic = True; continue
        if tok in ("</em>",): italic = False; continue
        if tok == "<b>": bold = True; continue
        if tok == "</b>": bold = False; continue
        if not tok: continue
        r = par.add_run(tok); r.bold = bold; r.italic = italic

i, n = 0, len(block)
while i < n:
    s = block[i].strip()
    if not s:
        i += 1; continue
    if s.startswith("# "):
        doc.add_heading(s[2:].strip(), level=1); i += 1; continue
    if s.startswith("## "):
        doc.add_heading(s[3:].strip(), level=2); i += 1; continue
    if s.startswith("### "):
        doc.add_heading(s[4:].strip(), level=3); i += 1; continue
    if s.startswith(">"):
        p = doc.add_paragraph(style="Quote"); add_inline(p, s.lstrip("> ").strip())
        i += 1; continue
    if s.startswith("- "):
        p = doc.add_paragraph(style="List Bullet"); add_inline(p, s[2:].strip())
        i += 1; continue
    if s.startswith('<div class="fiction">'):
        inner = re.sub(r'^<div class="fiction">|</div>\s*$', "", s)
        p = doc.add_paragraph(); p.paragraph_format.left_indent = Inches(0.35)
        p.paragraph_format.space_before = Pt(4); p.paragraph_format.space_after = Pt(8)
        m_tit = re.match(r'<span class="titolo-scena">(.*?)</span>(.*)$', inner)
        m_chi = re.match(r'<span class="chi">(.*?)</span>(.*)$', inner)
        if m_tit:
            r = p.add_run(m_tit.group(1)); r.bold = True; r.font.size = Pt(13)
            r.add_break(); add_inline(p, m_tit.group(2), italic=True)
        elif m_chi:
            r = p.add_run(m_chi.group(1) + ". "); r.bold = True; r.italic = True
            add_inline(p, m_chi.group(2))
        else:
            add_inline(p, inner)
        i += 1; continue
    # paragrafo normale; se la riga seguente e' il marcatore {: .azione } e' una riga d'azione
    is_azione = (i + 1 < n and block[i + 1].strip() == "{: .azione }")
    p = doc.add_paragraph()
    if is_azione:
        p.paragraph_format.space_before = Pt(10)
    add_inline(p, s)
    i += 2 if is_azione else 1

doc.save(OUT)
print("scritto", OUT)
